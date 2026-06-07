"""
agent/tools.py
--------------
All LangChain tools available to the Analyst AI ReAct agent.

FIXES APPLIED:
  1. search_documents tool: previously did `from rag.retriever import
     search_documents as _search` INSIDE the function body on every call.
     This imported a fresh function reference each time and bypassed the
     module-level _STORE cache in retriever.py. Fixed by importing the
     rag.retriever MODULE once at call time and using module.function(),
     which always reads the live _STORE variable.
  2. Same import pattern fixed in compare_sources, summarise_topic,
     extract_key_claims, and _build_fallback_content for consistency.
  3. _retry_with_backoff preserved for transient API failure handling.
  4. Removed risky _STORE = None cache-reset fallback in search_documents —
     an empty result does not mean the store is corrupted, and resetting
     caused unnecessary disk reloads on valid but empty queries.
"""

import os
import json
import re
import time
from datetime import datetime
from pathlib import Path

from langchain_core.tools import tool

# ─── Constants ────────────────────────────────────────────────────────────────
MAX_SEARCH_RESULTS  = 5
MAX_CONTENT_LENGTH  = 400
MAX_CSV_ROWS        = 500
SEARCH_DEPTH        = "advanced"
CACHE_TTL_MINUTES   = 60
DEFAULT_SESSION_ID  = 1
MIN_REPORT_WORDS    = 800
UPLOADS_DIR = Path(__file__).parent.parent / "uploads"
REPORTS_DIR = Path(__file__).parent.parent / "reports"

# ─── Retry Configuration ──────────────────────────────────────────────────────
MAX_RETRIES      = 3
RETRY_BASE_DELAY = 1.0
RETRY_BACKOFF    = 2.0

# ─── Runtime Context ──────────────────────────────────────────────────────────
# Set by app.py before each run_agent() call so tools can scope DB writes
# to the correct user without needing session state imports everywhere.
_current_user_id: int | None = None
_current_session_id: int = DEFAULT_SESSION_ID


def set_agent_context(user_id: int | None, session_id: int) -> None:
    """Set the user and session context for the current agent invocation.

    Called by app.py immediately before run_agent() so that all tool
    DB writes (reports, citations) are scoped to the correct user/session.

    Args:
        user_id: The logged-in user's database ID.
        session_id: The active session's database ID.
    """
    global _current_user_id, _current_session_id
    _current_user_id    = user_id
    _current_session_id = session_id


# ─── Retry Helper ─────────────────────────────────────────────────────────────

def _retry_with_backoff(fn, *args, retries: int = MAX_RETRIES, **kwargs):
    """Call fn with exponential back-off retry on transient failures.

    Args:
        fn: Callable to invoke.
        *args: Positional arguments forwarded to fn.
        retries: Maximum number of attempts.
        **kwargs: Keyword arguments forwarded to fn.

    Returns:
        The return value of fn on success.

    Raises:
        The last exception raised by fn if all retries are exhausted.
    """
    delay = RETRY_BASE_DELAY
    last_exc: Exception | None = None

    for attempt in range(1, retries + 1):
        try:
            return fn(*args, **kwargs)
        except (KeyboardInterrupt, SystemExit):
            raise
        except Exception as exc:
            last_exc = exc
            if attempt < retries:
                time.sleep(delay)
                delay *= RETRY_BACKOFF

    raise last_exc  # type: ignore[misc]


# ─── RAG Module Helper ────────────────────────────────────────────────────────

def _get_retriever_module():
    """Import and return the rag.retriever module.

    FIX: All tools import the MODULE (not individual functions) so they
    always call through the module's live _STORE cache. Importing the
    function directly (e.g. `from rag.retriever import search_documents`)
    captures the function object at import time and can miss cache updates
    that happen after the module was first loaded.

    Returns:
        The rag.retriever module object.
    """
    import rag.retriever as _retriever
    return _retriever


# ─── Document Search ──────────────────────────────────────────────────────────

@tool
def search_documents(query: str, source_filter: str | None = None) -> str:
    """Semantically search through all uploaded documents using FAISS vector search.

    Always use this tool first before web_search for document-based questions.
    Combines semantic similarity with optional source filtering.
    Retries up to MAX_RETRIES times on transient embedding-API failures.

    Args:
        query: The search query string to find relevant passages.
        source_filter: Optional filename to restrict search to a specific document.

    Returns:
        Formatted string of relevant passages with source citations and page numbers,
        or a message indicating no results were found.
    """
    # FIX: use module reference so we always hit the live _STORE cache
    retriever = _get_retriever_module()

    def _do_search():
        return retriever.search_documents(
            query, k=MAX_SEARCH_RESULTS, source_filter=source_filter
        )

    try:
        results = _retry_with_backoff(_do_search)
    except (FileNotFoundError, RuntimeError) as e:
        return f"❌ Search failed after {MAX_RETRIES} attempts: {str(e)}"
    except Exception as e:
        return f"❌ Search failed: {str(e)}"

    # FIX: Removed the risky _STORE = None cache-reset fallback that was here.
    # An empty result does not mean the store is corrupted — resetting _STORE
    # on an empty query caused unnecessary disk reloads and could mask valid
    # (but simply unmatched) queries. If results are empty, we report honestly.
    if not results:
        return (
            "📄 No relevant passages found for this query in the indexed documents. "
            "The documents ARE indexed — try rephrasing the query or use web_search "
            "for additional context. Do NOT ask the user to re-upload files."
        )

    output = [f"📄 **Document search results for:** '{query}'\n"]
    for i, r in enumerate(results, 1):
        page_info = f", p.{r['page']}" if r.get("page") else ""
        output.append(f"**[{i}] Source: {r['source']}{page_info}**")
        output.append(r["content"].strip())
        output.append("")
    return "\n".join(output)


# ─── Web Search ───────────────────────────────────────────────────────────────

@tool
def web_search(query: str) -> str:
    """Search the live web for current information using the Tavily search API.

    Checks a local cache before making an API call. Results are cached for
    CACHE_TTL_MINUTES to reduce redundant API usage.
    Retries up to MAX_RETRIES times with exponential back-off on failure.

    Use for: recent news, data not in uploaded documents, market info, current events.

    Args:
        query: The search query string.

    Returns:
        Formatted string of web results including titles, URLs, and content snippets,
        or an error message if the API key is missing or the request fails.
    """
    from db.database import cache_search, get_cached_search

    cached = get_cached_search(query)
    if cached:
        return _format_web_results(query, cached)

    api_key = os.getenv("TAVILY_API_KEY", "")
    if not api_key:
        return (
            "⚠️ Tavily API key not configured. "
            "Add TAVILY_API_KEY to your .env file. "
            "Get a free key at https://tavily.com"
        )

    def _do_tavily_search():
        from tavily import TavilyClient
        client = TavilyClient(api_key=api_key)
        return client.search(
            query=query,
            search_depth=SEARCH_DEPTH,
            max_results=MAX_SEARCH_RESULTS,
            include_answer=True,
        )

    try:
        response = _retry_with_backoff(_do_tavily_search)
        results = [
            {
                "title":   r.get("title", ""),
                "url":     r.get("url", ""),
                "content": r.get("content", ""),
                "score":   r.get("score", 0),
            }
            for r in response.get("results", [])
        ]
        cache_search(query, results)
        return _format_web_results(query, results, answer=response.get("answer"))

    except ImportError as e:
        return f"❌ Tavily package not installed: {str(e)}"
    except ValueError as e:
        return f"❌ Invalid Tavily API key: {str(e)}"
    except ConnectionError as e:
        return f"❌ Network error during web search after {MAX_RETRIES} attempts: {str(e)}"
    except Exception as e:
        return f"❌ Web search failed after {MAX_RETRIES} attempts: {str(e)}"


def _format_web_results(
    query: str, results: list, answer: str | None = None
) -> str:
    """Format raw Tavily results into a readable markdown string."""
    lines = [f"🌐 **Web search results for:** '{query}'\n"]
    if answer:
        lines.append(f"**Quick Answer:** {answer}\n")
    for i, r in enumerate(results, 1):
        content   = r["content"]
        truncated = content[:MAX_CONTENT_LENGTH].strip()
        if len(content) > MAX_CONTENT_LENGTH:
            truncated += "..."
        lines.append(f"**[{i}] {r['title']}**")
        lines.append(f"🔗 {r['url']}")
        lines.append(truncated)
        lines.append("")
    return "\n".join(lines)


# ─── Compare Sources ──────────────────────────────────────────────────────────

@tool
def compare_sources(topic: str, source_names: list | None = None) -> str:
    """Compare how multiple documents address a given topic side by side.

    Retrieves relevant chunks from each source and presents them together
    to highlight agreements, contradictions, and unique perspectives.

    Args:
        topic: The topic or question to compare across sources.
        source_names: Optional list of document filenames to compare.
            If None, all indexed documents are used.

    Returns:
        Formatted markdown string with per-source content and a synthesis note.
    """
    retriever = _get_retriever_module()

    try:
        sources = source_names or retriever.get_document_sources()
    except (FileNotFoundError, RuntimeError) as e:
        return f"❌ Could not retrieve document sources: {str(e)}"

    if not sources:
        return "⚠️ No documents in knowledge base. Please upload documents first."

    lines       = [f"## 📊 Source Comparison: '{topic}'\n"]
    all_content = {}

    for source in sources:
        try:
            results = _retry_with_backoff(
                retriever.search_documents, topic, k=3, source_filter=source
            )
            if results:
                combined = " ".join(r["content"] for r in results)[:800]
                all_content[source] = combined
        except (FileNotFoundError, RuntimeError):
            continue

    if not all_content:
        return f"No relevant content found across documents for topic: '{topic}'"

    lines.append("### What each source says:\n")
    for source, content in all_content.items():
        lines.append(f"**📄 {source}:**")
        lines.append(content.strip())
        lines.append("")

    lines.append("---")
    lines.append(
        f"*Ask the agent to analyse agreements and contradictions across "
        f"these sources on '{topic}'*"
    )
    return "\n".join(lines)


# ─── Summarise Topic ──────────────────────────────────────────────────────────

@tool
def summarise_topic(topic: str, include_web: bool = True) -> str:
    """Generate a comprehensive summary of a topic from uploaded documents.

    Args:
        topic: The topic to summarise.
        include_web: If True, prompts the agent to also run a web search.

    Returns:
        Structured markdown summary with document citations.
    """
    retriever = _get_retriever_module()
    lines     = [f"## 📋 Summary: {topic}\n"]

    if retriever.has_knowledge_base():
        try:
            doc_results = _retry_with_backoff(
                retriever.search_documents, topic, k=6
            )
            if doc_results:
                lines.append("### From Uploaded Documents:\n")
                for r in doc_results:
                    src  = r["source"]
                    page = f", p.{r['page']}" if r.get("page") else ""
                    lines.append(
                        f"**[Doc: {src}{page}]** "
                        f"{r['content'][:MAX_CONTENT_LENGTH].strip()}"
                    )
                    lines.append("")
        except (FileNotFoundError, RuntimeError) as e:
            lines.append(f"*Could not retrieve document content: {str(e)}*\n")

    if include_web:
        lines.append("### From Web Research:\n")
        lines.append(
            "*(Use web_search tool to add live web context to this summary)*"
        )

    lines.append("\n---")
    lines.append(
        "**Next steps:** Use `generate_report` to turn this into a downloadable "
        "report, or `compare_sources` to see how different sources differ."
    )
    return "\n".join(lines)


# ─── CSV Analysis ─────────────────────────────────────────────────────────────

@tool
def analyse_csv(filename: str, question: str | None = None) -> str:
    """Perform statistical analysis on an uploaded CSV file.

    Args:
        filename: The name of the CSV file in the uploads directory.
        question: Optional specific question about the data to address.

    Returns:
        Formatted markdown string with shape, column info, and numeric summary.
    """
    csv_path = UPLOADS_DIR / filename

    if not csv_path.exists():
        matches = list(UPLOADS_DIR.glob(f"*{filename}*"))
        if matches:
            csv_path = matches[0]
        else:
            available = [f.name for f in UPLOADS_DIR.iterdir() if f.is_file()]
            return f"❌ File '{filename}' not found. Available files: {available}"

    try:
        import pandas as pd
        df = pd.read_csv(str(csv_path))
    except (FileNotFoundError, PermissionError) as e:
        return f"❌ Could not open file: {str(e)}"
    except ValueError as e:
        return f"❌ Could not parse CSV: {str(e)}"
    except Exception as e:
        return f"❌ Could not analyse CSV: {str(e)}"

    lines = [f"## 📊 Analysis: {csv_path.name}\n"]
    lines.append(f"**Shape:** {df.shape[0]:,} rows × {df.shape[1]} columns\n")

    lines.append("**Columns:**")
    for col in df.columns:
        dtype    = str(df[col].dtype)
        null_pct = round(df[col].isna().mean() * 100, 1)
        lines.append(f"- `{col}` ({dtype}) — {null_pct}% missing")

    lines.append("\n**Numeric Summary:**")
    numeric_cols = df.select_dtypes(include="number")
    if not numeric_cols.empty:
        desc = numeric_cols.describe().round(2)
        lines.append("```")
        lines.append(desc.to_string())
        lines.append("```")

    cat_cols = df.select_dtypes(include=["object", "category"]).columns[:3]
    for col in cat_cols:
        top = df[col].value_counts().head(5)
        lines.append(f"\n**Top values in `{col}`:**")
        for val, count in top.items():
            lines.append(f"  - {val}: {count:,}")

    if question:
        lines.append(f"\n**Regarding your question:** '{question}'")
        lines.append(
            "*(The above statistics should help answer this — "
            "ask for a deeper analysis if needed)*"
        )

    return "\n".join(lines)


# ─── Extract Key Claims ───────────────────────────────────────────────────────

@tool
def extract_key_claims(
    source_name: str | None = None,
    max_claims: int = 10,
) -> str:
    """Extract the most important claims, facts, and conclusions from documents.

    Args:
        source_name: Optional filename to focus on one document.
        max_claims: Maximum number of claim passages to return.

    Returns:
        Formatted markdown list of key claims with source citations.
    """
    retriever   = _get_retriever_module()
    CLAIMS_QUERY = "key findings conclusions recommendations data evidence"

    try:
        if source_name:
            results = _retry_with_backoff(
                retriever.search_documents,
                CLAIMS_QUERY, k=max_claims, source_filter=source_name,
            )
            title = f"Key Claims from: {source_name}"
        else:
            results = _retry_with_backoff(
                retriever.search_documents, CLAIMS_QUERY, k=max_claims
            )
            title = "Key Claims Across All Documents"
    except (FileNotFoundError, RuntimeError) as e:
        return f"❌ Could not search documents: {str(e)}"

    if not results:
        return "📄 No documents found. Please upload documents first."

    lines = [f"## 🔍 {title}\n"]
    lines.append(
        "The following are the most relevant passages for key claims "
        "and findings:\n"
    )
    for i, r in enumerate(results, 1):
        page = f", p.{r['page']}" if r.get("page") else ""
        lines.append(f"**[{i}] [{r['source']}{page}]**")
        lines.append(r["content"][:500].strip())
        lines.append("")

    lines.append("---")
    lines.append(
        "*Ask the agent to synthesise these claims or identify the "
        "strongest evidence.*"
    )
    return "\n".join(lines)


# ─── Generate Report ──────────────────────────────────────────────────────────

@tool
def generate_report(
    title: str,
    content: str,
    report_type: str = "research",
    include_toc: bool = True,
) -> str:
    """Generate a professional downloadable .docx research report.

    Args:
        title: The report title shown on the cover page.
        content: Full report content in markdown format (800+ words ideal).
        report_type: One of 'research', 'summary', 'comparison', 'analysis'.
        include_toc: Reserved for future table-of-contents support.

    Returns:
        Success message with download instructions, or error with traceback.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        return f"❌ python-docx not installed: {str(e)}"

    try:
        REPORTS_DIR.mkdir(exist_ok=True)

        timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]
        filename   = f"{safe_title}_{timestamp}.docx"
        filepath   = REPORTS_DIR / filename

        if not content or len(content.strip()) < MIN_REPORT_WORDS // 4:
            content = _build_fallback_content(title)

        doc     = Document()
        section = doc.sections[0]
        for attr, val in [
            ("page_height",    Inches(11)),
            ("page_width",     Inches(8.5)),
            ("left_margin",    Inches(1)),
            ("right_margin",   Inches(1)),
            ("top_margin",     Inches(1)),
            ("bottom_margin",  Inches(1)),
        ]:
            setattr(section, attr, val)

        # Cover page
        title_para           = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run            = title_para.add_run(title)
        title_run.bold       = True
        title_run.font.size  = Pt(24)
        title_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        doc.add_paragraph()

        meta_para           = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run            = meta_para.add_run(
            f"Report Type: {report_type.title()}  |  "
            f"Generated: {datetime.now().strftime('%B %d, %Y')}\n"
            f"Generated by Analyst AI — Evidence-Driven Research Agent"
        )
        meta_run.font.size  = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        doc.add_page_break()

        # Content
        for line in content.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith("### "):
                h = doc.add_heading(line[4:], level=2)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            elif line.startswith("## "):
                h = doc.add_heading(line[3:], level=1)
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            elif line.startswith("# "):
                h = doc.add_heading(line[2:], level=1)
                if h.runs:
                    h.runs[0].bold = True
            elif line.startswith(("- ", "* ")):
                p = doc.add_paragraph(line[2:], style="List Bullet")
                if p.runs:
                    p.runs[0].font.size = Pt(11)
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                p    = doc.add_paragraph(text, style="List Number")
                if p.runs:
                    p.runs[0].font.size = Pt(11)
            elif line.startswith("---"):
                p   = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr   = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'CCCCCC')
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Inches(0.5)
                if p.runs:
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            else:
                p = doc.add_paragraph()
                for part in re.split(r'(\*\*[^*]+\*\*)', line):
                    if part.startswith("**") and part.endswith("**"):
                        run      = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.size = Pt(11)

        # Footer
        footer      = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run  = footer_para.add_run(
            f"Analyst AI — Research Agent  |  "
            f"{datetime.now().strftime('%B %d, %Y')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        doc.save(str(filepath))

        from db.database import save_report
        save_report(
            title=title,
            file_path=str(filepath),
            report_type=report_type,
            user_id=_current_user_id,
        )

        return (
            f"✅ **Report generated successfully!**\n\n"
            f"📄 **{title}**\n"
            f"🗓️ Created: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n"
            f"Go to the **📑 Reports** tab to download your report."
        )

    except OSError as e:
        return f"❌ File system error: {str(e)}"
    except Exception as e:
        import traceback
        return f"❌ Report generation failed: {str(e)}\n\n{traceback.format_exc()}"


def _build_fallback_content(title: str) -> str:
    """Build fallback report content from the knowledge base."""
    try:
        retriever = _get_retriever_module()
        queries   = [
            "key findings main topics overview",
            "history background introduction",
            "technology advancements developments",
            "conclusions summary recommendations",
        ]
        kb_content = []
        seen: set  = set()

        for q in queries:
            results = _retry_with_backoff(retriever.search_documents, q, k=4)
            for r in results:
                chunk = r["content"].strip()
                if chunk and chunk[:50] not in seen:
                    seen.add(chunk[:50])
                    src  = r.get("source", "document")
                    page = f", p.{r['page']}" if r.get("page") else ""
                    kb_content.append(f"{chunk}\n[Source: {src}{page}]")

        if kb_content:
            return (
                f"## Executive Summary\n\n"
                f"This report presents a comprehensive analysis based on "
                f"the uploaded research documents.\n\n"
                f"## Key Findings\n\n"
                + "\n".join(f"- {c[:400]}" for c in kb_content[:6])
                + f"\n\n## Detailed Analysis\n\n"
                + "\n\n".join(kb_content[6:12])
                + f"\n\n## Conclusions\n\n"
                f"Based on the analysis of the uploaded documents, the research "
                f"provides valuable insights.\n\n"
                f"## References\n\n"
                f"- Source documents indexed in the Analyst AI knowledge base.\n"
                f"- Generated on {datetime.now().strftime('%B %d, %Y')} by Analyst AI."
            )
    except (FileNotFoundError, RuntimeError):
        pass

    return (
        f"## {title}\n\n"
        f"Report generated on {datetime.now().strftime('%B %d, %Y')}.\n\n"
        f"No document content was available at the time of generation.\n"
        f"Please ensure documents are uploaded and ingested before generating a report."
    )


# Citation Tools 

@tool
def save_citation(
    title: str,
    source_type: str,
    authors: str = "",
    year: str = "",
    url: str = "",
    publisher: str = "",
    page_numbers: str = "",
    session_id: int = DEFAULT_SESSION_ID,
) -> str:
    """Save a source citation to the citation manager database.

    Args:
        title: Title of the source being cited.
        source_type: One of 'document', 'web', 'book', 'journal', 'report'.
        authors: Author name(s) as a string.
        year: Publication year as a string.
        url: Web URL if the source is online.
        publisher: Publisher or website name.
        page_numbers: Page numbers if citing a document.
        session_id: Database session ID to associate the citation with.

    Returns:
        Confirmation message with citation ID.
    """
    from db.database import save_citation as db_save_citation
    try:
        _sid = session_id if session_id != DEFAULT_SESSION_ID else _current_session_id
        cid = db_save_citation(
            session_id=_sid,
            source_type=source_type,
            title=title,
            authors=authors,
            year=year,
            url=url,
            publisher=publisher,
            page_numbers=page_numbers,
            citation_format="APA",
        )
    except Exception as e:
        return f"❌ Could not save citation: {str(e)}"

    return (
        f"📚 Citation saved (ID: {cid})\n"
        f"**{title}** — {source_type} source"
        + (f" by {authors}" if authors else "")
        + (f" ({year})" if year else "")
        + "\nView all citations in the **📚 Citations** tab."
    )


@tool
def list_citations(session_id: int = DEFAULT_SESSION_ID) -> str:
    """List all citations saved in the current or specified session.

    Args:
        session_id: Database session ID to retrieve citations for.

    Returns:
        Formatted markdown list of all saved citations.
    """
    from db.database import get_session_citations, format_citation
    try:
        _sid = session_id if session_id != DEFAULT_SESSION_ID else _current_session_id
        citations = get_session_citations(_sid)
    except Exception as e:
        return f"❌ Could not retrieve citations: {str(e)}"

    if not citations:
        return "📚 No citations saved yet. I'll automatically save citations as we research."

    lines      = [f"## 📚 Citations ({len(citations)} sources)\n"]
    type_icons = {
        "web": "🌐", "document": "📄", "book": "📕",
        "journal": "📰", "report": "📊",
    }
    for i, c in enumerate(citations, 1):
        formatted = format_citation(c)
        icon      = type_icons.get(c["source_type"], "📄")
        lines.append(f"**[{i}]** {icon} {formatted}")
        lines.append("")

    lines.append("---")
    lines.append(
        "*Go to the **📚 Citations** tab to export in APA, MLA, or Chicago format.*"
    )
    return "\n".join(lines)


# Tool Registry 

ALL_TOOLS = {
    "search_documents":  search_documents,
    "web_search":        web_search,
    "compare_sources":   compare_sources,
    "summarise_topic":   summarise_topic,
    "analyse_csv":       analyse_csv,
    "extract_key_claims": extract_key_claims,
    "generate_report":   generate_report,
    "save_citation":     save_citation,
    "list_citations":    list_citations,
}

TOOL_DESCRIPTIONS = {
    "search_documents":  "Search uploaded documents (RAG)",
    "web_search":        "Live web search (Tavily)",
    "compare_sources":   "Compare multiple sources",
    "summarise_topic":   "Summarise a topic",
    "analyse_csv":       "Analyse CSV/data files",
    "extract_key_claims": "Extract key claims & facts",
    "generate_report":   "Generate downloadable report",
    "save_citation":     "Save citations",
    "list_citations":    "List all citations",
}

CORE_TOOLS = list(ALL_TOOLS.keys())